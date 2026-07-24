import os
from docx import Document


OUTPUT_FOLDER = "generated_docs"

os.makedirs(
    OUTPUT_FOLDER,
    exist_ok=True
)


def generate_document(
    document_type: str,
    template: str,
    title: str,
    name: str,
    company: str,
    content: str
):
    """
    Generate a document based on the selected template.
    """

    if document_type.lower() != "docx":
        raise ValueError(
            "Currently only DOCX format is supported."
        )

    document = Document()

    # Resume Template
    if template.lower() == "resume":

        document.add_heading(
            "Resume",
            level=1
        )

        document.add_heading(
            "Name",
            level=2
        )
        document.add_paragraph(name)

        document.add_heading(
            "Skills",
            level=2
        )
        document.add_paragraph(content)

    # Cover Letter Template
    elif template.lower() == "cover_letter":

        document.add_heading(
            "Cover Letter",
            level=1
        )

        document.add_heading(
            "Applicant",
            level=2
        )
        document.add_paragraph(name)

        document.add_heading(
            "Company",
            level=2
        )
        document.add_paragraph(company)

        document.add_heading(
            "Letter",
            level=2
        )
        document.add_paragraph(content)

    # Statement of Purpose Template
    elif template.lower() == "sop":

        document.add_heading(
            "Statement of Purpose",
            level=1
        )

        document.add_heading(
            "Applicant",
            level=2
        )
        document.add_paragraph(name)

        document.add_heading(
            "Purpose",
            level=2
        )
        document.add_paragraph(content)

    # Project Report Template
    elif template.lower() == "project_report":

        document.add_heading(
            "Project Report",
            level=1
        )

        document.add_heading(
            "Project Title",
            level=2
        )
        document.add_paragraph(title)

        document.add_heading(
            "Prepared By",
            level=2
        )
        document.add_paragraph(name)

        document.add_heading(
            "Report",
            level=2
        )
        document.add_paragraph(content)

    # Default Template
    else:

        document.add_heading(
            title,
            level=1
        )

        document.add_paragraph(content)

    filename = (
        title.replace(" ", "_")
        + ".docx"
    )

    filepath = os.path.join(
        OUTPUT_FOLDER,
        filename
    )

    document.save(filepath)

    return {
        "message": "Document generated successfully.",
        "filename": filename
    }